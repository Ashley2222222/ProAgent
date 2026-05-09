# -*- coding: utf-8 -*-
import json
import os
import io
import threading
import uuid
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from urllib.parse import parse_qs, urlparse, quote, unquote

import requests
from urllib3.exceptions import InsecureRequestWarning

from ai_analyzer import ResearchAIAnalyzer
from crawler_factory import ResearchCrawlerFactory
from html_report_generator import generate_html_report
from proposal_generator import generate_full_proposal, generate_proposals_batch
from rag_analyzer import analyze_materials_batch
from knowledge_base_volc import KnowledgeBase
from config import OUTPUT_DIR, RAG_RESULT_FILENAME, RESULT_FILENAME, WECHAT_WEBHOOK_URL

requests.packages.urllib3.disable_warnings(InsecureRequestWarning)
from config import (
    OUTPUT_DIR,
    RAG_RESULT_FILENAME,
    RESULT_FILENAME,
    WECHAT_WEBHOOK_URL,
    PROJECT_ROOT,
)

JOBS = {}
JOBS_LOCK = threading.Lock()
KB_LOCK = threading.Lock()
KB_INSTANCE = None
NOTICE_LOCK = threading.Lock()
NOTICE_INDEX = {}
NOTICE_LIST = []


def _now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _log(job_id: str, message: str) -> None:
    line = f"[{_now_str()}] {message}"
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if job is not None:
            job["logs"].append(line)


def _send_wechat_message(message: str) -> str:
    if not WECHAT_WEBHOOK_URL:
        return "未配置企业微信 WECHAT_WEBHOOK_URL"
    headers = {"Content-Type": "application/json"}
    data = {"msgtype": "markdown", "markdown": {"content": message}}
    resp = requests.post(
        WECHAT_WEBHOOK_URL, headers=headers, json=data, verify=False, timeout=30
    )
    resp.raise_for_status()
    r = resp.json()
    if r.get("errcode") == 0:
        return "推送成功"
    return f"推送失败: {r.get('errmsg')}"


def _extract_wechat_webhook_key(webhook_url: str) -> str | None:
    try:
        if not webhook_url:
            return None
        if "key=" not in webhook_url:
            return None
        return webhook_url.split("key=", 1)[1].split("&", 1)[0].strip() or None
    except Exception:
        return None


def _upload_file_to_wechat(file_path: str, webhook_key: str) -> str | None:
    try:
        if not file_path or not os.path.exists(file_path):
            return None
        file_size = os.path.getsize(file_path)
        if not (5 <= file_size <= 20 * 1024 * 1024):
            return None
        if not webhook_key:
            return None

        upload_url = f"https://qyapi.weixin.qq.com/cgi-bin/webhook/upload_media?key={webhook_key}&type=file"
        with open(file_path, "rb") as f:
            files = {
                "media": (os.path.basename(file_path), f, "application/octet-stream")
            }
            resp = requests.post(upload_url, files=files, timeout=60, verify=False)
        resp.raise_for_status()
        result = resp.json()
        if result.get("errcode") == 0 and result.get("media_id"):
            return result["media_id"]
        return None
    except Exception:
        return None


def _send_wechat_file(media_id: str) -> str:
    if not WECHAT_WEBHOOK_URL:
        return "未配置企业微信 WECHAT_WEBHOOK_URL"
    if not media_id:
        return "media_id 为空"
    headers = {"Content-Type": "application/json"}
    data = {"msgtype": "file", "file": {"media_id": media_id}}
    resp = requests.post(
        WECHAT_WEBHOOK_URL, headers=headers, json=data, verify=False, timeout=30
    )
    resp.raise_for_status()
    r = resp.json()
    if r.get("errcode") == 0:
        return "文件推送成功"
    return f"文件推送失败: {r.get('errmsg')}"


def _parse_bool(v):
    if isinstance(v, bool):
        return v
    if isinstance(v, str):
        return v.lower() in ("1", "true", "yes", "on")
    return False


def _get_kb() -> KnowledgeBase:
    global KB_INSTANCE
    with KB_LOCK:
        if KB_INSTANCE is None:
            kb_dir = os.path.join(OUTPUT_DIR, "kb")
            KB_INSTANCE = KnowledgeBase(persist_dir=kb_dir)
        return KB_INSTANCE


def _make_notice_id(item: dict) -> str:
    link = (item.get("链接") or "").strip()
    title = (item.get("标题") or "").strip()
    base = link or title
    return str(abs(hash(base)))


def _refresh_notices_from_disk() -> None:
    candidates = [
        os.path.join(OUTPUT_DIR, "analysis_results.json"),
        os.path.join(OUTPUT_DIR, RESULT_FILENAME),
        os.path.join(OUTPUT_DIR, "crawl_results.json"),
    ]
    data = None
    for p in candidates:
        if os.path.exists(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    data = json.load(f)
                break
            except Exception:
                data = None
    if not isinstance(data, list):
        return

    idx = {}
    lst = []
    for item in data:
        if not isinstance(item, dict):
            continue
        nid = _make_notice_id(item)
        idx[nid] = item
        lst.append(
            {
                "id": nid,
                "title": item.get("标题", ""),
                "site": item.get("网站", ""),
                "date": item.get("发布日期", ""),
            }
        )
    with NOTICE_LOCK:
        NOTICE_INDEX.clear()
        NOTICE_INDEX.update(idx)
        NOTICE_LIST.clear()
        NOTICE_LIST.extend(lst)


def _md_to_docx_bytes(markdown_text: str) -> bytes:
    from docx import Document

    doc = Document()
    lines = (markdown_text or "").replace("\r\n", "\n").split("\n")
    for line in lines:
        t = line.rstrip()
        if not t.strip():
            doc.add_paragraph("")
            continue
        if t.startswith("### "):
            doc.add_heading(t[4:].strip(), level=3)
            continue
        if t.startswith("## "):
            doc.add_heading(t[3:].strip(), level=2)
            continue
        if t.startswith("# "):
            doc.add_heading(t[2:].strip(), level=1)
            continue
        if t.startswith("- "):
            doc.add_paragraph(t[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(t)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _run_pipeline(job_id: str, options: dict) -> None:
    with JOBS_LOCK:
        JOBS[job_id]["status"] = "running"
        JOBS[job_id]["options"] = options

    enable_ai = _parse_bool(options.get("enable_ai"))
    enable_html = _parse_bool(options.get("enable_html"))
    enable_wechat_markdown = _parse_bool(
        options.get("enable_wechat_markdown", options.get("enable_wechat"))
    )
    enable_wechat_file = _parse_bool(options.get("enable_wechat_file"))
    enable_proposal = _parse_bool(options.get("enable_proposal"))
    start_date = (options.get("start_date") or "").strip() or None
    end_date = (options.get("end_date") or "").strip() or None
    selected_sites = options.get("sites") or []
    if isinstance(selected_sites, str):
        selected_sites = [selected_sites]

    site_configs = ResearchCrawlerFactory.get_all_sites()
    if not selected_sites:
        selected_sites = list(site_configs.keys())
    selected_sites = [s for s in selected_sites if s in site_configs]

    _log(
        job_id,
        f"开始执行：网站 {len(selected_sites)} 个，AI分析={enable_ai}，HTML={enable_html}，企业微信摘要={enable_wechat_markdown}，企业微信文件={enable_wechat_file}",
    )
    analyzer = ResearchAIAnalyzer() if enable_ai else None

    all_results = []

    try:
        for site_key in selected_sites:
            config = site_configs[site_key]
            _log(job_id, f"网站：{config.get('description')} ({site_key})")
            crawler = ResearchCrawlerFactory.create_crawler(site_key)

            all_notices = []
            for list_config in config.get("list_urls", []):
                list_url = list_config["url"]
                list_name = list_config.get("name", "未命名")
                _log(job_id, f"列表页：{list_name} {list_url}")
                max_pages = config.get("max_pages", 3)
                notices = crawler.fetch_notice_list(list_url, max_pages=max_pages) or []
                _log(job_id, f"抓取列表：{len(notices)} 条")
                all_notices.extend(notices)

            pre_filter_count = len(all_notices)
            all_notices = [
                n
                for n in all_notices
                if crawler.should_include_notice(n.get("标题", ""))
            ]
            if pre_filter_count != len(all_notices):
                _log(
                    job_id,
                    f"标题过滤：保留 {len(all_notices)} 条（剔除 {pre_filter_count - len(all_notices)} 条）",
                )

            deduped = []
            seen = set()
            for n in all_notices:
                link = (n.get("链接") or "").strip()
                title = (n.get("标题") or "").strip()
                key = link or title
                if not key:
                    continue
                if key in seen:
                    continue
                seen.add(key)
                deduped.append(n)
            if len(deduped) != len(all_notices):
                _log(
                    job_id,
                    f"去重：保留 {len(deduped)} 条（剔除 {len(all_notices) - len(deduped)} 条重复）",
                )
            all_notices = deduped

            for idx, notice in enumerate(all_notices, 1):
                title = notice.get("标题") or ""
                link = notice.get("链接") or ""
                _log(job_id, f"详情抓取：{idx}/{len(all_notices)} {title[:40]}")
                detail = crawler.fetch_notice_detail(link)
                content = detail.get("content", "") if isinstance(detail, dict) else ""
                pub_date = notice.get("发布日期")
                if isinstance(detail, dict) and detail.get("pub_date"):
                    pub_date = detail.get("pub_date")

                if start_date or end_date:
                    if not crawler.is_date_in_range(pub_date, start_date, end_date):
                        continue

                if not content or "失败" in content:
                    continue

                if enable_ai and analyzer is not None:
                    ai_result = analyzer.analyze_notice(title, content)
                    result = {
                        "网站": crawler.get_site_name(),
                        "标题": title,
                        "发布日期": pub_date,
                        "链接": link,
                        "正文长度": len(content),
                        "正文预览": content[:30000],
                        "分析结果": ai_result,
                    }
                else:
                    result = {
                        "网站": crawler.get_site_name(),
                        "标题": title,
                        "发布日期": pub_date,
                        "链接": link,
                        "正文长度": len(content),
                        "正文预览": content[:30000],
                    }
                all_results.append(result)

        os.makedirs(os.path.join(PROJECT_ROOT, OUTPUT_DIR), exist_ok=True)
        crawl_path = os.path.join(PROJECT_ROOT, OUTPUT_DIR, RESULT_FILENAME)
        with open(crawl_path, "w", encoding="utf-8") as f:
            json.dump(all_results, f, ensure_ascii=False, indent=2)
        _log(job_id, f"爬取结果已保存：{crawl_path}")
        _refresh_notices_from_disk()

        results_for_html = all_results
        rag_path = None
        html_path = None
        html_filename = None
        materials_results = []
        if enable_ai and all_results:
            _log(job_id, "开始材料提取（基于AI分析结果）")
            materials_results = analyze_materials_batch(all_results)
            rag_path = os.path.join(PROJECT_ROOT, OUTPUT_DIR, RAG_RESULT_FILENAME)
            with open(rag_path, "w", encoding="utf-8") as f:
                json.dump(materials_results, f, ensure_ascii=False, indent=2)
            _log(job_id, f"材料提取结果已保存:{rag_path}")

        if enable_proposal and all_results:
            proposals_dir = os.path.join(OUTPUT_DIR, "proposals")
            if enable_ai and materials_results:
                _log(job_id, "开始生成申报建议书（基于材料提取结果）")
                count = generate_proposals_batch(
                    materials_results, output_dir=proposals_dir
                )
                _log(job_id, f"建议书生成完成：{count} 份，目录：{proposals_dir}")
            else:
                _log(job_id, "建议书生成已跳过：需先启用AI分析并产生材料提取结果")

        if enable_html:
            # 决定用于显示的数据源：如果启用了AI分析，优先使用 all_results（含分析结果）
            if enable_ai:
                display_results = all_results
            else:
                display_results = results_for_html if results_for_html else all_results
            if display_results:
                html_filename = (
                    f"政策爬取日报_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
                )
                data_dir = OUTPUT_DIR
                os.makedirs(data_dir, exist_ok=True)
                html_path = os.path.join(data_dir, html_filename)
                generate_html_report(display_results, html_path)
                _log(job_id, f"HTML日报已生成：{html_path}")

        if enable_wechat_markdown:
            _log(job_id, "开始企业微信推送（摘要）")
            today = datetime.now().strftime("%Y-%m-%d")
            msg_title = f"【科研课题申报日报】{today}"
            message = f"## {msg_title}\n\n"
            message += f"- **网站**: {len(selected_sites)} 个\n"
            message += f"- **通知**: {len(all_results)} 条\n"
            if start_date or end_date:
                message += (
                    f"- **日期范围**: {start_date or '不限'} 至 {end_date or '不限'}\n"
                )
            if html_filename:
                message += f"- **HTML日报**: {html_filename}\n"
            if rag_path:
                message += f"- **RAG结果**: {os.path.basename(rag_path)}\n"
            res = _send_wechat_message(message)
            _log(job_id, f"企业微信摘要：{res}")

        if enable_wechat_file:
            if not html_path or not os.path.exists(html_path):
                _log(job_id, "企业微信文件：未生成HTML或文件不存在，跳过发送")
            else:
                _log(job_id, "开始企业微信推送（HTML文件）")
                webhook_key = _extract_wechat_webhook_key(WECHAT_WEBHOOK_URL) or ""
                media_id = _upload_file_to_wechat(html_path, webhook_key)
                if not media_id:
                    _log(job_id, "企业微信文件：上传失败（未获取到 media_id）")
                else:
                    res = _send_wechat_file(media_id)
                    _log(job_id, f"企业微信文件：{res}")

        with JOBS_LOCK:
            JOBS[job_id]["status"] = "done"
            JOBS[job_id]["outputs"] = {
                "crawl_path": crawl_path,
                "rag_path": rag_path,
                "html_path": html_path,
                "html_filename": html_filename,
            }
    except Exception as e:
        _log(job_id, f"执行失败：{str(e)}")
        with JOBS_LOCK:
            JOBS[job_id]["status"] = "error"
            JOBS[job_id]["error"] = str(e)


def _run_generate_proposal(
    job_id: str, notice_id: str, top_k: int, save_to_kb: bool
) -> None:
    with JOBS_LOCK:
        JOBS[job_id]["status"] = "running"

    try:
        _log(job_id, "开始生成完整申报书（RAG知识库）")
        _refresh_notices_from_disk()
        with NOTICE_LOCK:
            notice = NOTICE_INDEX.get(notice_id)
        if not notice:
            raise ValueError("未找到通知，请先执行爬取生成通知列表")

        kb = _get_kb()
        if kb.is_empty():
            raise ValueError("知识库为空，请先上传历史申报书")

        md = generate_full_proposal(notice, kb, top_k=top_k)
        if save_to_kb and md and not md.strip().startswith("⚠️"):
            kb.add_text(
                text=md,
                source_name=f"generated_{notice_id}.md",
                metadata={"type": "generated", "notice_id": notice_id},
            )
            _log(job_id, "已将生成的申报书写入知识库")

        with JOBS_LOCK:
            JOBS[job_id]["status"] = "done"
            JOBS[job_id]["outputs"] = {
                "proposal_md": md,
                "notice_title": notice.get("标题", ""),
            }
    except Exception as e:
        _log(job_id, f"生成失败：{str(e)}")
        with JOBS_LOCK:
            JOBS[job_id]["status"] = "error"
            JOBS[job_id]["error"] = str(e)


def _html_page() -> str:
    sites = ResearchCrawlerFactory.get_all_sites()
    site_items = []
    for k, cfg in sites.items():
        desc = cfg.get("description") or k
        site_items.append(
            f'<label class="site"><input type="checkbox" name="sites" value="{k}" checked> {desc} <span class="muted">({k})</span></label>'
        )
    sites_html = "\n".join(site_items)
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>科研课题申报智能体</title>
  <style>
    body {{ font-family: -apple-system,BlinkMacSystemFont,Segoe UI,Roboto,Helvetica Neue,Arial,sans-serif; background:#f6f7fb; margin:0; }}
    .wrap {{ max-width: 1100px; margin: 24px auto; padding: 0 16px; }}
    .card {{ background:#fff; border-radius: 12px; box-shadow: 0 6px 22px rgba(0,0,0,.06); overflow:hidden; }}
    .header {{ padding: 16px 20px; background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); color:#fff; }}
    .content {{ padding: 18px 20px; display:grid; grid-template-columns: 1fr 1fr; gap: 16px; }}
    .box {{ border: 1px solid #eef0f6; border-radius: 10px; padding: 14px; }}
    .row {{ display:flex; gap:12px; flex-wrap:wrap; align-items:center; }}
    label {{ font-size: 14px; }}
    input[type="date"] {{ padding: 8px 10px; border: 1px solid #dfe3ee; border-radius: 8px; }}
    .checks {{ display:flex; gap:14px; flex-wrap:wrap; }}
    .site-list {{ max-height: 320px; overflow:auto; padding-right: 6px; }}
    .site {{ display:block; padding: 8px 10px; border-radius: 8px; border: 1px solid #eef0f6; margin-bottom: 8px; }}
    .muted {{ color:#6b7280; font-size:12px; }}
    .actions {{ padding: 0 20px 18px 20px; display:flex; gap:12px; }}
    button {{ border:0; padding: 10px 14px; border-radius: 10px; cursor:pointer; font-weight:600; }}
    .primary {{ background: #2563eb; color:#fff; }}
    .ghost {{ background:#eef2ff; color:#1f3a8a; }}
    .log {{ white-space: pre-wrap; background:#0b1220; color:#cbd5e1; border-radius: 12px; padding: 14px; font-family: ui-monospace,SFMono-Regular,Menlo,Consolas,monospace; font-size: 12px; max-height: 360px; overflow:auto; }}
    .footer {{ padding: 14px 20px; border-top: 1px solid #eef0f6; display:flex; justify-content:space-between; gap:10px; flex-wrap:wrap; }}
    a {{ color:#2563eb; text-decoration:none; }}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="card">
      <div class="header">
        <div style="font-size:18px;font-weight:700;">科研课题申报智能体（Web）</div>
        <div style="opacity:.95;margin-top:6px;font-size:13px;">选择日期与网站，点击执行。AI分析 / 生成HTML / 企业微信推送可选。</div>
      </div>
      <div class="content">
        <div class="box">
          <div style="font-weight:700;margin-bottom:10px;">日期范围</div>
          <div class="row">
            <label>开始 <input id="start_date" type="date"></label>
            <label>结束 <input id="end_date" type="date"></label>
          </div>
          <div style="font-weight:700;margin:14px 0 10px;">选项</div>
          <div class="checks">
            <label><input id="enable_ai" type="checkbox" checked> AI分析</label>
            <label><input id="enable_html" type="checkbox" checked> 生成HTML</label>
            <label><input id="enable_wechat_markdown" type="checkbox"> 推送摘要</label>
            <label><input id="enable_wechat_file" type="checkbox" checked> 推送HTML文件</label>
            <label><input id="enable_proposal" type="checkbox" checked> 生成建议书</label>
          </div>
        </div>
        <div class="box">
          <div style="font-weight:700;margin-bottom:10px;">网站选择</div>
          <div class="site-list" id="site_list">
            {sites_html}
          </div>
          <div class="row" style="margin-top:10px;">
            <button class="ghost" type="button" onclick="selectAll(true)">全选</button>
            <button class="ghost" type="button" onclick="selectAll(false)">全不选</button>
          </div>
        </div>
      </div>
      <div class="content" style="grid-template-columns: 1fr; padding-top: 0;">
        <div class="box">
          <div style="font-weight:700;margin-bottom:10px;">📚 历史申报书知识库</div>
          <div class="row" style="margin-bottom:10px;">
            <input id="kb_files" type="file" multiple accept=".docx,.md,.txt,.pdf">
            <button class="ghost" type="button" onclick="uploadKbFiles()">上传并建立知识库</button>
            <span class="muted" id="kb_status"></span>
          </div>
          <div class="muted">支持 .docx/.md/.txt/.pdf（单个≤100MB），向量库持久化在 data/kb</div>
        </div>
        <div class="box">
          <div style="font-weight:700;margin-bottom:10px;">✍️ 一键生成申报书</div>
          <div class="row" style="margin-bottom:10px;">
            <label>选择通知
              <select id="notice_select" style="padding:8px 10px;border:1px solid #dfe3ee;border-radius:8px;min-width:520px;"></select>
            </label>
            <label>TopK
              <input id="kb_topk" type="number" min="1" max="20" value="5" style="width:80px;padding:8px 10px;border:1px solid #dfe3ee;border-radius:8px;">
            </label>
            <label><input id="kb_save_back" type="checkbox"> 写入知识库</label>
            <button class="primary" type="button" onclick="startGenerateProposal()">智能生成申报书</button>
          </div>
          <div class="row" style="margin-bottom:10px;">
            <button class="ghost" type="button" onclick="downloadProposal('md')" id="btn_dl_md" disabled>下载MD</button>
            <button class="ghost" type="button" onclick="downloadProposal('docx')" id="btn_dl_docx" disabled>下载DOCX</button>
            <span class="muted" id="proposal_status"></span>
          </div>
          <div id="proposal_preview" style="border:1px solid #eef0f6;border-radius:10px;padding:12px;min-height:120px;background:#fbfcff;"></div>
        </div>
      </div>
      <div class="actions">
        <button class="primary" onclick="runJob()">开始执行</button>
        <button class="ghost" onclick="clearLog()">清空日志</button>
        <div id="result_links" style="align-self:center;"></div>
      </div>
      <div style="padding: 0 20px 18px 20px;">
        <div class="log" id="log"></div>
      </div>
      <div class="footer">
        <div class="muted">提示：执行期间可保持页面打开，日志会实时刷新。</div>
        <div class="muted">服务端仅在本机运行：建议用于内网/本机环境。</div>
      </div>
    </div>
  </div>

  <script>
    let currentJobId = null;
    let logOffset = 0;

    function selectAll(v) {{
      document.querySelectorAll('input[name=\"sites\"]').forEach(cb => cb.checked = v);
    }}

    function clearLog() {{
      document.getElementById('log').textContent = '';
      document.getElementById('result_links').innerHTML = '';
      logOffset = 0;
    }}

    function getOptions() {{
      const sites = Array.from(document.querySelectorAll('input[name=\"sites\"]:checked')).map(x => x.value);
      return {{
        start_date: document.getElementById('start_date').value,
        end_date: document.getElementById('end_date').value,
        sites,
        enable_ai: document.getElementById('enable_ai').checked,
        enable_html: document.getElementById('enable_html').checked,
        enable_wechat_markdown: document.getElementById('enable_wechat_markdown').checked,
        enable_wechat_file: document.getElementById('enable_wechat_file').checked,
        enable_proposal: document.getElementById('enable_proposal').checked
      }};
    }}

    async function runJob() {{
      clearLog();
      const options = getOptions();
      const resp = await fetch('/api/run', {{
        method: 'POST',
        headers: {{ 'Content-Type': 'application/json' }},
        body: JSON.stringify(options)
      }});
      const data = await resp.json();
      currentJobId = data.job_id;
      pollStatus();
    }}

    async function pollStatus() {{
      if (!currentJobId) return;
      const resp = await fetch('/api/status?id=' + encodeURIComponent(currentJobId));
      const data = await resp.json();
      const logs = data.logs || [];
      if (logs.length > logOffset) {{
        const newLines = logs.slice(logOffset).join('\\n') + '\\n';
        const el = document.getElementById('log');
        el.textContent += newLines;
        el.scrollTop = el.scrollHeight;
        logOffset = logs.length;
      }}
      if (data.outputs && data.outputs.html_filename) {{
        const url = '/files/' + encodeURIComponent(data.outputs.html_filename);
        document.getElementById('result_links').innerHTML = `<a href="${{url}}" target="_blank">打开HTML日报</a>`;
      }}
      if (data.status === 'running') {{
        setTimeout(pollStatus, 1200);
      }}
      if (data.status === 'error') {{
        setTimeout(pollStatus, 2000);
      }}
    }}

    let proposalJobId = null;
    let lastProposalJobId = null;

    function escapeHtml(s) {{
      return (s || '').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
    }}

    function renderMarkdown(md) {{
      const lines = (md || '').replace(/\\r\\n/g,'\\n').split('\\n');
      const out = [];
      for (const line of lines) {{
        if (line.startsWith('### ')) {{
          out.push('<h3>' + escapeHtml(line.slice(4)) + '</h3>');
        }} else if (line.startsWith('## ')) {{
          out.push('<h2>' + escapeHtml(line.slice(3)) + '</h2>');
        }} else if (line.startsWith('# ')) {{
          out.push('<h1>' + escapeHtml(line.slice(2)) + '</h1>');
        }} else if (line.startsWith('- ')) {{
          out.push('<div style="padding-left:18px">• ' + escapeHtml(line.slice(2)) + '</div>');
        }} else {{
          out.push('<div>' + escapeHtml(line) + '</div>');
        }}
      }}
      return out.join('');
    }}

    async function refreshNotices() {{
      const sel = document.getElementById('notice_select');
      if (!sel) return;
      sel.innerHTML = '<option value="">（加载中...）</option>';
      try {{
        const resp = await fetch('/api/notices');
        const data = await resp.json();
        const items = data.items || [];
        if (!items.length) {{
          sel.innerHTML = '<option value="">（暂无通知，请先执行爬取）</option>';
          return;
        }}
        sel.innerHTML = items.map(x => {{
          const t = (x.title || '').slice(0, 80);
          const d = x.date || '';
          const s = x.site || '';
          return `<option value="${{x.id}}">${{t}} | ${{s}} | ${{d}}</option>`;
        }}).join('');
      }} catch (e) {{
        sel.innerHTML = '<option value="">（加载失败）</option>';
      }}
    }}

    async function uploadKbFiles() {{
      const el = document.getElementById('kb_files');
      const st = document.getElementById('kb_status');
      if (!el || !el.files || el.files.length === 0) {{
        if (st) st.textContent = '请选择文件';
        return;
      }}
      if (st) st.textContent = '上传中...';
      const fd = new FormData();
      for (const f of el.files) fd.append('files', f, f.name);
      const resp = await fetch('/api/upload_proposal', {{ method: 'POST', body: fd }});
      const data = await resp.json();
      if (st) st.textContent = data.message || '完成';
    }}

    async function startGenerateProposal() {{
      const sel = document.getElementById('notice_select');
      const st = document.getElementById('proposal_status');
      if (!sel || !sel.value) {{
        if (st) st.textContent = '请选择通知';
        return;
      }}
      const topk = parseInt(document.getElementById('kb_topk').value || '5', 10);
      const saveBack = document.getElementById('kb_save_back').checked;
      document.getElementById('proposal_preview').innerHTML = '';
      document.getElementById('btn_dl_md').disabled = true;
      document.getElementById('btn_dl_docx').disabled = true;
      if (st) st.textContent = '生成中...';

      const resp = await fetch('/api/generate_proposal', {{
        method: 'POST',
        headers: {{ 'Content-Type': 'application/json' }},
        body: JSON.stringify({{ notice_id: sel.value, top_k: topk, save_to_kb: saveBack }})
      }});
      const data = await resp.json();
      proposalJobId = data.job_id;
      lastProposalJobId = proposalJobId;
      pollProposal();
    }}

    async function pollProposal() {{
      if (!proposalJobId) return;
      const st = document.getElementById('proposal_status');
      const resp = await fetch('/api/status?id=' + encodeURIComponent(proposalJobId));
      const data = await resp.json();
      if (data.status === 'running') {{
        if (st) st.textContent = '生成中...';
        setTimeout(pollProposal, 1200);
        return;
      }}
      if (data.status === 'error') {{
        if (st) st.textContent = '生成失败：' + (data.error || '');
        return;
      }}
      const md = data.outputs && data.outputs.proposal_md ? data.outputs.proposal_md : '';
      if (st) st.textContent = '完成';
      document.getElementById('proposal_preview').innerHTML = renderMarkdown(md);
      document.getElementById('btn_dl_md').disabled = !md;
      document.getElementById('btn_dl_docx').disabled = !md;
    }}

    function downloadProposal(fmt) {{
      const id = lastProposalJobId;
      if (!id) return;
      const url = '/api/download_proposal?id=' + encodeURIComponent(id) + '&format=' + encodeURIComponent(fmt);
      window.open(url, '_blank');
    }}

    window.addEventListener('load', () => {{
      refreshNotices();
    }});
  </script>
</body>
</html>"""


def _json_response(handler: BaseHTTPRequestHandler, code: int, obj: dict) -> None:
    data = json.dumps(obj, ensure_ascii=False).encode("utf-8")
    handler.send_response(code)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def _text_response(
    handler: BaseHTTPRequestHandler,
    code: int,
    text: str,
    content_type: str = "text/html; charset=utf-8",
) -> None:
    data = text.encode("utf-8")
    handler.send_response(code)
    handler.send_header("Content-Type", content_type)
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def _safe_report_path(filename: str) -> str | None:
    if not filename:
        return None
    if not filename.endswith(".html"):
        return None
    if not filename.startswith("政策爬取日报_"):
        return None
    data_dir = OUTPUT_DIR
    full = os.path.abspath(os.path.join(data_dir, filename))
    if not full.startswith(os.path.abspath(data_dir) + os.sep):
        return None
    if os.path.exists(full):
        return full
    full_root = os.path.abspath(os.path.join(PROJECT_ROOT, filename))
    if full_root.startswith(os.path.abspath(PROJECT_ROOT) + os.sep) and os.path.exists(
        full_root
    ):
        return full_root
    return None


class Handler(BaseHTTPRequestHandler):
    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path == "/":
            return _text_response(self, 200, _html_page())
        if parsed.path == "/api/notices":
            _refresh_notices_from_disk()
            with NOTICE_LOCK:
                items = list(NOTICE_LIST)
            return _json_response(self, 200, {"items": items})
        if parsed.path == "/api/status":
            q = parse_qs(parsed.query)
            job_id = (q.get("id") or [""])[0]
            with JOBS_LOCK:
                job = JOBS.get(job_id)
                if not job:
                    return _json_response(self, 404, {"error": "job_not_found"})
                return _json_response(
                    self,
                    200,
                    {
                        "status": job.get("status"),
                        "logs": job.get("logs", []),
                        "outputs": job.get("outputs", {}),
                        "error": job.get("error", ""),
                    },
                )
        if parsed.path == "/api/download_proposal":
            q = parse_qs(parsed.query)
            job_id = (q.get("id") or [""])[0]
            fmt = (q.get("format") or ["md"])[0].lower()
            with JOBS_LOCK:
                job = JOBS.get(job_id)
                outputs = (job or {}).get("outputs") or {}
            md = outputs.get("proposal_md") or ""
            title = outputs.get("notice_title") or "proposal"
            safe_name = (
                "".join(
                    c for c in str(title) if c.isalnum() or c in (" ", "-", "_")
                ).strip()[:50]
                or "proposal"
            )
            if not md:
                return _text_response(
                    self, 404, "not found", "text/plain; charset=utf-8"
                )
            if fmt == "docx":
                try:
                    data = _md_to_docx_bytes(md)
                except Exception as e:
                    return _text_response(
                        self, 500, str(e), "text/plain; charset=utf-8"
                    )
                self.send_response(200)
                self.send_header(
                    "Content-Type",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                safe_file_name = f"{safe_name}.docx"
                safe_file_name_quoted = quote(safe_file_name)
                self.send_header(
                    "Content-Disposition",
                    f"attachment; filename*=UTF-8''{safe_file_name_quoted}",
                )
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return
            # 默认 md 格式
            data = md.encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/markdown; charset=utf-8")
            safe_file_name_md = f"{safe_name}.md"
            safe_file_name_md_quoted = quote(safe_file_name_md)
            self.send_header(
                "Content-Disposition",
                f"attachment; filename*=UTF-8''{safe_file_name_md_quoted}",
            )
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return

        if parsed.path.startswith("/files/proposals/"):
            raw_filename = parsed.path[len("/files/proposals/") :]
            filename = unquote(raw_filename)
            if not filename.endswith(".md"):
                return _text_response(
                    self, 404, "not found", "text/plain; charset=utf-8"
                )
            proposals_dir = os.path.join(OUTPUT_DIR, "proposals")
            file_path = os.path.join(proposals_dir, filename)
            if os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "text/plain; charset=utf-8")
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return
            return _text_response(self, 404, "not found", "text/plain; charset=utf-8")
        if parsed.path.startswith("/files/"):
            raw_filename = parsed.path[len("/files/") :]
            filename = unquote(raw_filename)
            file_path = os.path.join(OUTPUT_DIR, filename)
            if os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "text/html; charset=utf-8")
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return
            else:
                file_path_root = os.path.join(PROJECT_ROOT, filename)
                if os.path.exists(file_path_root):
                    with open(file_path_root, "rb") as f:
                        data = f.read()
                    self.send_response(200)
                    self.send_header("Content-Type", "text/html; charset=utf-8")
                    self.send_header("Content-Length", str(len(data)))
                    self.end_headers()
                    self.wfile.write(data)
                    return
                return _text_response(
                    self, 404, "not found", "text/plain; charset=utf-8"
                )
        if parsed.path.startswith("/proposal/"):
            raw_filename = parsed.path[len("/proposal/") :]
            filename = unquote(raw_filename)
            if not filename.endswith(".md"):
                return _text_response(
                    self, 404, "not found", "text/plain; charset=utf-8"
                )
            proposals_dir = os.path.join(OUTPUT_DIR, "proposals")
            file_path = os.path.join(proposals_dir, filename)
            if os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "text/plain; charset=utf-8")
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return
            else:
                return _text_response(
                    self, 404, "not found", "text/plain; charset=utf-8"
                )

    def do_POST(self):
        parsed = urlparse(self.path)
        if parsed.path == "/api/run":
            length = int(self.headers.get("Content-Length") or "0")
            body = self.rfile.read(length) if length > 0 else b"{}"
            try:
                options = json.loads(body.decode("utf-8"))
            except Exception:
                options = {}
            job_id = uuid.uuid4().hex
            with JOBS_LOCK:
                JOBS[job_id] = {"status": "queued", "logs": [], "outputs": {}}
            _log(job_id, "任务已创建")
            t = threading.Thread(
                target=_run_pipeline, args=(job_id, options), daemon=True
            )
            t.start()
            return _json_response(self, 200, {"job_id": job_id})
        if parsed.path == "/api/upload_proposal":
            try:
                content_type = self.headers.get("Content-Type", "")
                if "multipart/form-data" not in content_type:
                    return _json_response(
                        self, 400, {"message": "invalid content-type"}
                    )

                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length)

                from email.parser import BytesParser
                from email import policy

                fake_headers = f"Content-Type: {content_type}\n\n".encode("utf-8")
                full_data = fake_headers + body
                msg = BytesParser(policy=policy.default).parsebytes(full_data)

                files_data = []
                for part in msg.iter_parts():
                    if part.get_content_disposition() == "form-data":
                        name = part.get_param("name", header="Content-Disposition")
                        if name == "files" and part.get_filename():
                            filename = part.get_filename()
                            file_bytes = part.get_payload(decode=True)
                            files_data.append((filename, file_bytes))

                if not files_data:
                    return _json_response(self, 400, {"message": "no files"})

                kb = _get_kb()
                upload_dir = os.path.join(OUTPUT_DIR, "kb", "uploads")
                os.makedirs(upload_dir, exist_ok=True)

                ok_files = 0
                total_chunks = 0
                for filename, data in files_data:
                    ext = os.path.splitext(filename)[1].lower()
                    if ext not in (".docx", ".md", ".txt", ".pdf"):
                        continue
                    if len(data) > 100 * 1024 * 1024:
                        continue
                    save_name = f"{uuid.uuid4().hex}_{filename}"
                    save_path = os.path.join(upload_dir, save_name)
                    with open(save_path, "wb") as f:
                        f.write(data)
                    chunks = kb.add_file(save_path, original_name=filename)
                    ok_files += 1
                    total_chunks += int(chunks or 0)

                return _json_response(
                    self,
                    200,
                    {
                        "message": f"上传完成：{ok_files} 个文件，入库切片 {total_chunks} 条",
                        "files": ok_files,
                        "chunks": total_chunks,
                    },
                )
            except Exception as e:
                import traceback

                traceback.print_exc()
                return _json_response(self, 500, {"message": str(e)})
        if parsed.path == "/api/generate_proposal":
            length = int(self.headers.get("Content-Length") or "0")
            body = self.rfile.read(length) if length > 0 else b"{}"
            try:
                payload = json.loads(body.decode("utf-8"))
            except Exception:
                payload = {}
            notice_id = str(payload.get("notice_id") or "").strip()
            top_k = int(payload.get("top_k") or 5)
            save_to_kb = _parse_bool(payload.get("save_to_kb"))
            if not notice_id:
                return _json_response(self, 400, {"message": "missing notice_id"})
            job_id = uuid.uuid4().hex
            with JOBS_LOCK:
                JOBS[job_id] = {
                    "status": "queued",
                    "logs": [],
                    "outputs": {},
                    "kind": "proposal",
                }
            _log(job_id, "任务已创建")
            t = threading.Thread(
                target=_run_generate_proposal,
                args=(job_id, notice_id, top_k, save_to_kb),
                daemon=True,
            )
            t.start()
            return _json_response(self, 200, {"job_id": job_id})
        return _json_response(self, 404, {"error": "not_found"})

    def log_message(self, format, *args):
        return


def serve(host: str = "127.0.0.1", port: int = 8000):
    httpd = ThreadingHTTPServer((host, port), Handler)
    print(f"Web已启动: http://{host}:{port}/", flush=True)
    httpd.serve_forever()


if __name__ == "__main__":
    serve()
