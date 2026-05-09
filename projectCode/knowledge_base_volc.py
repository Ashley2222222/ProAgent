# -*- coding: utf-8 -*-
import os
import uuid
import requests
import pytesseract
from pdf2image import convert_from_path
from typing import List, Dict, Any

# 设置 Tesseract 路径（根据你的实际安装路径调整）
pytesseract.pytesseract.tesseract_cmd = r"D:\Tesseract-OCR\tesseract.exe"


class KnowledgeBase:
    def __init__(self, persist_dir: str, collection_name: str = "proposal_kb"):
        self.persist_dir = persist_dir
        self.collection_name = collection_name
        os.makedirs(self.persist_dir, exist_ok=True)
        self._client = None
        self._collection = None
        self._splitter = None

        # 火山引擎配置（从环境变量读取）
        self.api_key = os.getenv("AI_API_KEY")
        if not self.api_key:
            raise ValueError("请在环境变量中设置 AI_API_KEY")
        # 使用你创建的接入点 ID
        self.endpoint_id = "ep-20260511100953-rnlb2"  # 👈 替换为你的接入点 ID
        self.embedding_url = "https://ark.cn-beijing.volces.com/api/v3/embeddings"

    def _get_embedding(self, text: str) -> List[float]:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        input_content = [{"type": "text", "text": text[:4000]}]
        payload = {
            "model": self.endpoint_id,
            "input": input_content,
        }
        url = "https://ark.cn-beijing.volces.com/api/v3/embeddings/multimodal"
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            # ✅ 修复：根据真实返回结构取值
            embedding = data.get("data", {}).get("embedding")
            if not embedding:
                raise ValueError("响应中缺少 embedding 字段")
            return embedding
        except Exception as e:
            print(f"Embedding API 调用失败: {e}")
            raise

    def _ensure_ready(self) -> None:
        if self._client is None:
            import chromadb

            self._client = chromadb.PersistentClient(path=self.persist_dir)
        if self._collection is None:
            self._collection = self._client.get_or_create_collection(
                name=self.collection_name
            )
        if self._splitter is None:
            from langchain_text_splitters import RecursiveCharacterTextSplitter

            self._splitter = RecursiveCharacterTextSplitter(
                chunk_size=800, chunk_overlap=100
            )

    def is_empty(self) -> bool:
        self._ensure_ready()
        try:
            return self._collection.count() <= 0
        except Exception:
            return True

    def add_text(
        self,
        text: str,
        source_name: str,
        doc_id: str | None = None,
        metadata: dict | None = None,
    ) -> int:
        """添加单段文本到知识库（自动切片）"""
        if not text or not text.strip():
            return 0
        self._ensure_ready()

        doc_id = doc_id or uuid.uuid4().hex
        meta_base = {"source": source_name, "doc_id": doc_id}
        if isinstance(metadata, dict):
            meta_base.update(metadata)

        # 文本切片
        chunks = self._splitter.split_text(text)
        if not chunks:
            return 0

        # 为每个切片生成向量
        ids = []
        embeddings = []
        documents = []
        metadatas = []
        for i, chunk in enumerate(chunks):
            vec = self._get_embedding(chunk)
            ids.append(f"{doc_id}_{i}")
            embeddings.append(vec)
            documents.append(chunk)
            m = dict(meta_base)
            m["chunk_index"] = i
            metadatas.append(m)

        # 批量 upsert
        self._collection.upsert(
            ids=ids, embeddings=embeddings, documents=documents, metadatas=metadatas
        )
        return len(chunks)

    def add_file(
        self,
        file_path: str,
        original_name: str | None = None,
        metadata: dict | None = None,
    ) -> int:
        """上传文件（支持 .docx, .md, .txt, .pdf）"""
        if not file_path or not os.path.exists(file_path):
            return 0
        name = original_name or os.path.basename(file_path)
        ext = os.path.splitext(name)[1].lower()

        text = ""
        if ext == ".docx":
            from docx import Document

            doc = Document(file_path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        elif ext in (".md", ".txt"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif ext == ".pdf":
            import pdfplumber

            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            # OCR 回退（如果提取文字少于100字符）
            if len(text.strip()) < 100:
                try:
                    images = convert_from_path(
                        file_path,
                        dpi=200,
                        poppler_path=r"D:\poppler-26.02.0\Library\bin",  # 根据你的实际路径调整
                    )
                    ocr_text = ""
                    for img in images:
                        page_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                        ocr_text += page_text + "\n"
                    text = ocr_text
                except Exception as e:
                    print(f"OCR 处理失败: {e}")
                    text = ""
        else:
            return 0

        if not text.strip():
            return 0
        return self.add_text(text=text, source_name=name, metadata=metadata)

    def query(self, query_text: str, top_k: int = 5) -> list[dict]:
        """检索相似文本"""
        if not query_text or not query_text.strip():
            return []
        self._ensure_ready()
        vec = self._get_embedding(query_text)
        res = self._collection.query(
            query_embeddings=[vec],
            n_results=top_k,
            include=["documents", "metadatas", "distances"],
        )
        out = []
        docs = (res.get("documents") or [[]])[0]
        metas = (res.get("metadatas") or [[]])[0]
        dists = (res.get("distances") or [[]])[0]
        for i, doc in enumerate(docs):
            out.append(
                {
                    "text": doc,
                    "metadata": metas[i] if i < len(metas) else {},
                    "distance": dists[i] if i < len(dists) else None,
                }
            )
        return out
